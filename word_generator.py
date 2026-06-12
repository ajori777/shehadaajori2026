import os
import tempfile
from docxtpl import DocxTemplate
from docx.table import Table
import win32com.client
from PyPDF2 import PdfWriter, PdfReader

class WordGenerator:
    def __init__(self, template_path="m.docx"):
        self.template_path = template_path

    def fill_student_docx(self, student_data, school_info, output_path):
        tpl = DocxTemplate(self.template_path)
        
        avg_pct = str(student_data.get('المعدل_المئوي', '')).replace('%', '')
        result = str(student_data.get('النتيجة', '')).strip()
        
        if not result and avg_pct:
            try:
                avg_val = float(avg_pct.replace(',', '.'))
                if avg_val >= 50: result = "ناجح"
                else: result = "مقصر"
            except: pass

        # Absence logic
        total_abs = 0
        try:
            total_abs = int(float(str(student_data.get('الغياب', '0')).strip() or 0))
        except: pass
        
        abs1_val = total_abs // 2
        abs2_val = (total_abs // 2) + (total_abs % 2)
        
        abs1 = str(abs1_val) if abs1_val > 0 else "---"
        abs2 = str(abs2_val) if abs2_val > 0 else "---"

        def sanitize(v):
            if v is None or str(v).lower() == 'none':
                return ""
            return str(v).strip()

        def format_year(year_str):
            if not year_str: return ""
            # Extract numbers
            nums = __import__('re').findall(r'\d+', str(year_str))
            if len(nums) >= 2:
                # Sort descending to put higher year first
                sorted_nums = sorted([int(n) for n in nums], reverse=True)
                return f"{sorted_nums[0]} - {sorted_nums[1]}"
            return str(year_str)

        # Field/Stream logic for ST8
        st8_grade = sanitize(student_data.get('الصف', ''))
        st8_field = sanitize(student_data.get('field', ''))
        st8_section = sanitize(student_data.get('الشعبة', ''))
        if st8_field:
            st8_val = f"{st8_grade} / {st8_field} / {st8_section}"
        else:
            st8_val = f"{st8_grade} / {st8_section}"

        # Religion logic (ST6)
        religion = ""
        is_f1 = student_data.get('التربية الإسلامية_ف1')
        is_f2 = student_data.get('التربية الإسلامية_ف2')
        if (is_f1 and str(is_f1).strip()) or (is_f2 and str(is_f2).strip()):
            religion = "الإسلام"

        context = {
            'YEAR1': format_year(school_info.get('year', '')),
            'CLAS1': sanitize(school_info.get('grade', '')),
            'ST1': sanitize(student_data.get('name', '')),
            'ST2': sanitize(student_data.get('national_id', '')),
            'ST3': sanitize(student_data.get('dob', '')),
            'ST4': sanitize(student_data.get('pob', '')),
            'ST5': sanitize(student_data.get('nationality', '')),
            'ST6': religion,
            'ST8': st8_val,
            'SC1': sanitize(school_info.get('school_name', '')),
            'SC2': sanitize(school_info.get('national_id', '')),
            'SC6': sanitize(school_info.get('directorate', '')),
            'SC7': sanitize(school_info.get('district', '')),
            'SC3': sanitize(school_info.get('town', '')),
            'AVG': sanitize(student_data.get('المعدل_المئوي', '')),
            'RESULT': sanitize(result),
            'ABS1': sanitize(abs1),
            'ABS2': sanitize(abs2),
            'ABSTOTAL': sanitize(school_info.get('actual_days', '')),
            'teacher': sanitize(student_data.get('مربي_الصف', '')),
            'moder': sanitize(student_data.get('مدير_المدرسة', school_info.get('principal_name', ''))),
            'mar': sanitize(student_data.get('mar', ''))
        }
        
        tpl.render(context)
        doc = tpl
        
        # Fill subjects table via XPath
        all_tbl_elements = doc._element.xpath('//w:tbl')
        subjects_table = None
        for tbl_el in all_tbl_elements:
            table = Table(tbl_el, doc)
            if len(table.rows) > 0:
                header_text = "".join([c.text for c in table.rows[0].cells])
                if "المبحث" in header_text:
                    subjects_table = table
                    break
        
        if subjects_table:
            subjects = []
            seen_subjects = set()
            for key in student_data.keys():
                if key.endswith("_ف1"):
                    s_name = key[:-3] # Remove '_ف1'
                    if s_name not in seen_subjects:
                        subjects.append(s_name)
                        seen_subjects.add(s_name)
            
            # Determine how many rows to fill (starting from row 2)
            # Available rows from row 2 onwards
            available_rows = len(subjects_table.rows) - 2
            needed_rows = len(subjects)
            
            # Adjust table rows to match the number of subjects
            if needed_rows > available_rows:
                # Add rows and copy formatting from the last available row
                for _ in range(needed_rows - available_rows):
                    last_row = subjects_table.rows[-1]
                    new_row = subjects_table.add_row()
                    new_row.height = last_row.height
                    for i_cell, cell in enumerate(last_row.cells):
                        if i_cell < len(new_row.cells):
                            new_row.cells[i_cell]._tc.get_or_add_tcPr().append(
                                __import__('copy').deepcopy(cell._tc.get_or_add_tcPr())
                            )
            elif needed_rows < available_rows:
                # Delete extra rows from the bottom
                # We start from row index 2, so total rows should be needed_rows + 2
                while len(subjects_table.rows) > needed_rows + 2:
                    row_to_remove = subjects_table.rows[-1]
                    row_to_remove._element.getparent().remove(row_to_remove._element)

            for i in range(len(subjects)):
                row_idx = i + 2
                row = subjects_table.rows[row_idx]
                
                # To preserve formatting, we clear text of first run instead of cell.text
                def set_cell_text(cell, text):
                    if text is None: text = ""
                    if not cell.paragraphs:
                        cell.text = str(text)
                        return
                    # Clear all paragraphs but keep first one's formatting
                    p = cell.paragraphs[0]
                    p.text = str(text)
                    # Remove extra paragraphs if any
                    for extra_p in cell.paragraphs[1:]:
                        p_element = extra_p._element
                        p_element.getparent().remove(p_element)

                # Clear placeholder or existing data
                for cell in row.cells:
                    set_cell_text(cell, "")
                
                s_name = subjects[i]
                try:
                    set_cell_text(row.cells[5], s_name)
                    # Use the extracted max/min grades
                    max_val = student_data.get(f"{s_name}_ف1_عظمى", student_data.get(f"{s_name}_معدل_عظمى", "100"))
                    min_val = student_data.get(f"{s_name}_ف1_صغرى", student_data.get(f"{s_name}_معدل_صغرى", "50"))
                    
                    set_cell_text(row.cells[4], sanitize(min_val))
                    set_cell_text(row.cells[3], sanitize(max_val))
                    set_cell_text(row.cells[2], sanitize(student_data.get(f"{s_name}_ف1", "")))
                    set_cell_text(row.cells[1], sanitize(student_data.get(f"{s_name}_ف2", "")))
                    set_cell_text(row.cells[0], sanitize(student_data.get(f"{s_name}_معدل", "")))
                except: pass
        
        doc.save(output_path)

    def generate_section_file(self, students, school_info, pdf_path=None, word_path=None, progress_callback=None, limit=None):
        if not students: return
        
        # Apply trial limit if provided
        if limit and len(students) > limit:
            students = students[:limit]
            
        temp_dir = tempfile.mkdtemp()
        docx_paths = []
        
        # 1. Fill all docx templates first
        total = len(students)
        for i, student in enumerate(students):
            docx_path = os.path.join(temp_dir, f"student_{i}.docx")
            self.fill_student_docx(student, school_info, docx_path)
            docx_paths.append(docx_path)
            if progress_callback:
                progress_callback(i + 1, total)

        # 2. Use Word to merge and export
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        word.ScreenUpdating = False
        
        main_doc = None
        try:
            # Open the first document as a base to preserve margins and layout settings
            main_doc = word.Documents.Open(os.path.abspath(docx_paths[0]))
            word.Options.Pagination = True
            
            if len(docx_paths) > 1:
                sel = word.Selection
                for i in range(1, len(docx_paths)):
                    # Move to end of document
                    sel.EndKey(Unit=6) # wdStory
                    # Insert page break
                    sel.InsertBreak(Type=7) # wdPageBreak
                    # Insert next file
                    sel.InsertFile(os.path.abspath(docx_paths[i]))
            
            # Force update of layout
            main_doc.Repaginate()

            # Export as PDF if requested
            if pdf_path:
                abs_pdf = os.path.abspath(pdf_path)
                if not abs_pdf.lower().endswith('.pdf'):
                    abs_pdf += ".pdf"
                try:
                    # SaveAs2 is often more stable for Arabic layout in merged docs
                    main_doc.SaveAs2(abs_pdf, FileFormat=17) # wdFormatPDF
                except Exception as e:
                    raise Exception(f"خطأ أثناء حفظ ملف PDF: {str(e)}")
            
            # Save as Word (DOCX) if requested
            if word_path:
                abs_word = os.path.abspath(word_path)
                if not abs_word.lower().endswith('.docx'):
                    abs_word += ".docx"
                try:
                    main_doc.SaveAs2(abs_word, FileFormat=16) # wdFormatXMLDocument
                except Exception as e:
                    raise Exception(f"خطأ أثناء حفظ ملف Word: {str(e)}")
            
            main_doc.Close(SaveChanges=False)
            main_doc = None
                
        except Exception as e:
            if main_doc:
                try: main_doc.Close(SaveChanges=False)
                except: pass
            raise e
        finally:
            word.Quit()
            # Cleanup temp files
            for root, dirs, files in os.walk(temp_dir, topdown=False):
                for name in files:
                    try: os.remove(os.path.join(root, name))
                    except: pass
                for name in dirs:
                    try: os.rmdir(os.path.join(root, name))
                    except: pass
            try: os.rmdir(temp_dir)
            except: pass
